#!/usr/bin/env python3
# -*- coding: utf-8 -*-

"""
Script para generar documento Word con documentación técnica completa
del Portal Web del Instituto de Investigaciones Sociales
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from datetime import datetime

def add_heading_style(doc, text, level=1):
    """Agregar encabezado con estilo"""
    heading = doc.add_heading(text, level=level)
    return heading

def add_paragraph_style(doc, text, bold=False, italic=False, color=None):
    """Agregar párrafo con estilo"""
    p = doc.add_paragraph(text)
    if bold or italic or color:
        for run in p.runs:
            if bold:
                run.bold = True
            if italic:
                run.italic = True
            if color:
                run.font.color.rgb = color
    return p

def shade_cell(cell, color):
    """Sombrear celda en tabla"""
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), color)
    cell._element.get_or_add_tcPr().append(shading_elm)

def create_documentation():
    """Crear documento Word completo"""
    
    # Crear documento
    doc = Document()
    
    # Márgenes
    sections = doc.sections
    for section in sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)
    
    # Portada
    title = doc.add_heading('Portal Web del Instituto de Investigaciones Sociales', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subtitle = doc.add_paragraph('Documentación Técnica Completa')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].bold = True
    subtitle.runs[0].font.size = Pt(16)
    
    doc.add_paragraph()  # Espaciador
    
    info = [
        f'Documento de Especificación y Guía de Instalación',
        f'Fecha: Diciembre 2025',
        f'Versión: 1.0',
        f'Autor: Equipo de Desarrollo IIS',
        f'Institución: Instituto de Investigaciones Sociales, UNAM'
    ]
    
    for text in info:
        p = doc.add_paragraph(text)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_page_break()
    
    # Tabla de contenidos (manual)
    add_heading_style(doc, 'Tabla de Contenidos', 1)
    toc_items = [
        '1. Introducción',
        '2. Descripción General del Proyecto',
        '3. Arquitectura del Sistema',
        '4. Tecnologías Utilizadas',
        '5. Componentes del Sistema',
        '6. Requisitos de Instalación',
        '7. Guía de Instalación',
        '8. Configuración y Despliegue',
        '9. API Reference',
        '10. Estructura del Código',
        '11. Mantenimiento y Troubleshooting'
    ]
    
    for item in toc_items:
        doc.add_paragraph(item, style='List Bullet')
    
    doc.add_page_break()
    
    # 1. INTRODUCCIÓN
    add_heading_style(doc, '1. Introducción', 1)
    
    intro_text = """El Portal Web del Instituto de Investigaciones Sociales es una plataforma web integrada diseñada para facilitar el acceso, búsqueda y visualización de recursos digitales y contenido editorial.

El portal combina dos fuentes principales de contenido:
• Omeka 2.x: Repositorio digital para recursos académicos y culturales
• CMS Simple (Node.js): Sistema de gestión de contenidos para artículos y documentación editorial

Este documento proporciona la documentación técnica completa del prototipo, incluyendo arquitectura, componentes, requisitos de instalación y guías de despliegue."""
    
    doc.add_paragraph(intro_text)
    
    doc.add_page_break()
    
    # 2. DESCRIPCIÓN GENERAL
    add_heading_style(doc, '2. Descripción General del Proyecto', 1)
    
    add_heading_style(doc, '2.1 Objetivo', 2)
    objetivos = [
        'Catalogar y organizar recursos digitales del IIS',
        'Proporcionar una búsqueda unificada en múltiples fuentes',
        'Ofrecer acceso fácil a contenido académico y editorial',
        'Mantener una interfaz intuitiva y responsive'
    ]
    for obj in objetivos:
        doc.add_paragraph(obj, style='List Bullet')
    
    add_heading_style(doc, '2.2 Características Principales', 2)
    features = [
        'Búsqueda Unificada: Busca simultáneamente en Omeka (repositorio) y CMS (artículos)',
        'Interfaz Moderna: Diseño responsive con Vue.js 3',
        'API REST: Endpoints para integración con sistemas externos',
        'CORS Habilitado: Permite acceso desde diferentes dominios',
        'Almacenamiento Flexible: Combina bases de datos SQL y JSON',
        'Docker Ready: Despliegue simple con Docker Compose'
    ]
    for feat in features:
        doc.add_paragraph('✓ ' + feat, style='List Bullet')
    
    add_heading_style(doc, '2.3 Usuarios Objetivo', 2)
    users = [
        'Investigadores y académicos del IIS',
        'Público general interesado en recursos del Instituto',
        'Administradores del portal (para gestión de contenido)'
    ]
    for user in users:
        doc.add_paragraph(user, style='List Bullet')
    
    doc.add_page_break()
    
    # 3. ARQUITECTURA
    add_heading_style(doc, '3. Arquitectura del Sistema', 1)
    
    add_heading_style(doc, '3.1 Descripción de Capas', 2)
    
    doc.add_paragraph('La arquitectura se divide en tres capas principales:')
    
    layers = [
        ('Capa de Presentación', 'Vue.js 3 en el navegador. Interfaz responsiva y moderna.'),
        ('Capa de APIs', 'Omeka (repositorio), CMS Simple (artículos), búsqueda unificada.'),
        ('Capa de Datos', 'MySQL (Omeka), JSON (CMS Simple), almacenamiento persistente.')
    ]
    
    for layer_name, desc in layers:
        p = doc.add_paragraph()
        p.add_run(layer_name).bold = True
        p.add_run(': ' + desc)
    
    add_heading_style(doc, '3.2 Flujo de Comunicación', 2)
    
    flow_text = """1. Cliente (navegador) envía solicitud HTTP
2. Nginx (reverse proxy) recibe la solicitud
3. Vite (dev server) o Nginx sirve el frontend Vue.js
4. Vue.js ejecuta en el navegador
5. Cuando el usuario busca, Vue.js llama a la API
6. Los servicios backend (Omeka, CMS) procesan la solicitud
7. Los datos se envían de vuelta al navegador en JSON
8. Vue.js actualiza la interfaz reactivamente"""
    
    doc.add_paragraph(flow_text)
    
    doc.add_page_break()
    
    # 4. TECNOLOGÍAS
    add_heading_style(doc, '4. Tecnologías Utilizadas', 1)
    
    add_heading_style(doc, '4.1 Frontend', 2)
    
    table = doc.add_table(rows=7, cols=3)
    table.style = 'Light Grid Accent 1'
    
    header_cells = table.rows[0].cells
    headers = ['Tecnología', 'Versión', 'Propósito']
    for i, header in enumerate(headers):
        header_cells[i].text = header
        shade_cell(header_cells[i], '4472C4')
        for paragraph in header_cells[i].paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
                run.font.color.rgb = RGBColor(255, 255, 255)
    
    frontend_data = [
        ('Vue.js', '3.x', 'Framework reactivo'),
        ('Vite', '4.x', 'Build tool moderno'),
        ('Axios', '1.x', 'Cliente HTTP'),
        ('Vue Router', '4.x', 'Enrutamiento de páginas'),
        ('Pinia', '2.x', 'Gestión de estado'),
        ('DOMPurify', '3.x', 'Sanitización HTML')
    ]
    
    for i, (tech, version, purpose) in enumerate(frontend_data, 1):
        row_cells = table.rows[i].cells
        row_cells[0].text = tech
        row_cells[1].text = version
        row_cells[2].text = purpose
    
    add_heading_style(doc, '4.2 Backend', 2)
    
    table2 = doc.add_table(rows=6, cols=3)
    table2.style = 'Light Grid Accent 1'
    
    header_cells2 = table2.rows[0].cells
    for i, header in enumerate(headers):
        header_cells2[i].text = header
        shade_cell(header_cells2[i], '70AD47')
        for paragraph in header_cells2[i].paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
                run.font.color.rgb = RGBColor(255, 255, 255)
    
    backend_data = [
        ('Omeka (Repositorio)', 'PHP 7.4 / Apache 2.4', 'Gestión de recursos'),
        ('Base Datos Omeka', 'MySQL 8.0', 'Almacenamiento persistente'),
        ('CMS Simple', 'Node.js 22 / Express 4.18', 'Gestión de artículos'),
        ('Almacenamiento CMS', 'JSON (archivo)', 'Base de datos simple'),
        ('Reverse Proxy', 'Nginx Alpine', 'Enrutamiento de solicitudes')
    ]
    
    for i, (tech, version, purpose) in enumerate(backend_data, 1):
        row_cells = table2.rows[i].cells
        row_cells[0].text = tech
        row_cells[1].text = version
        row_cells[2].text = purpose
    
    doc.add_page_break()
    
    # 5. COMPONENTES
    add_heading_style(doc, '5. Componentes del Sistema', 1)
    
    add_heading_style(doc, '5.1 Portal Web (Frontend)', 2)
    
    doc.add_paragraph('Ubicación: /src/')
    
    doc.add_paragraph('Componentes principales:', style='List Bullet')
    for comp in ['App.vue: Componente raíz', 'Header.vue: Encabezado y navegación', 
                 'Footer.vue: Pie de página', 'Navigation.vue: Menú de navegación']:
        doc.add_paragraph(comp, style='List Bullet 2')
    
    doc.add_paragraph('Vistas principales:', style='List Bullet')
    for view in ['Home.vue: Página de inicio', 'Search.vue: Búsqueda avanzada',
                 'Resources.vue: Catálogo de recursos', 'ResourceDetail.vue: Detalle de recurso']:
        doc.add_paragraph(view, style='List Bullet 2')
    
    add_heading_style(doc, '5.2 Omeka 2.x (Repositorio Digital)', 2)
    
    omeka_desc = """Omeka es un framework de código abierto para crear colecciones de objetos digitales y compartirlas en línea.

Características:
• Gestión centralizada de recursos digitales
• API REST para integración con aplicaciones externas
• Soporte para metadatos complejos (Dublin Core)
• CORS habilitado para acceso desde navegadores
• MySQL como base de datos persistente"""
    
    doc.add_paragraph(omeka_desc)
    
    add_heading_style(doc, '5.3 CMS Simple (Node.js)', 2)
    
    cms_desc = """Sistema de gestión de contenidos lightweight basado en Node.js y Express.

Características:
• API REST completa (CRUD)
• Almacenamiento en JSON (sin dependencias externas)
• Búsqueda integrada
• UTF-8 para caracteres especiales
• Fácil de extender y personalizar"""
    
    doc.add_paragraph(cms_desc)
    
    doc.add_page_break()
    
    # 6. REQUISITOS
    add_heading_style(doc, '6. Requisitos de Instalación', 1)
    
    add_heading_style(doc, '6.1 Requisitos de Hardware', 2)
    
    hw_reqs = [
        ('CPU', '2+ cores (recomendado 4)'),
        ('RAM', '4GB mínimo (8GB recomendado)'),
        ('Almacenamiento', '20GB libre (SSD recomendado)'),
        ('Conexión', 'Acceso a internet para descargas')
    ]
    
    for req, spec in hw_reqs:
        p = doc.add_paragraph()
        p.add_run(req + ': ').bold = True
        p.add_run(spec)
    
    add_heading_style(doc, '6.2 Requisitos de Software', 2)
    
    add_heading_style(doc, 'Windows', 3)
    doc.add_paragraph('Docker Desktop for Windows (versión 4.0+)', style='List Bullet')
    doc.add_paragraph('Git for Windows (versión 2.30+)', style='List Bullet')
    doc.add_paragraph('PowerShell 5.1 (incluido en Windows 10/11)', style='List Bullet')
    
    add_heading_style(doc, 'macOS', 3)
    doc.add_paragraph('Docker Desktop for Mac (versión 4.0+)', style='List Bullet')
    doc.add_paragraph('Git (incluido o brew install git)', style='List Bullet')
    
    add_heading_style(doc, 'Linux', 3)
    doc.add_paragraph('Docker CE (versión 20.10+)', style='List Bullet')
    doc.add_paragraph('Docker Compose (versión 2.0+)', style='List Bullet')
    doc.add_paragraph('Git', style='List Bullet')
    
    add_heading_style(doc, '6.3 Puertos Requeridos', 2)
    
    table3 = doc.add_table(rows=8, cols=4)
    table3.style = 'Light Grid Accent 1'
    
    header_cells3 = table3.rows[0].cells
    port_headers = ['Servicio', 'Puerto', 'Protocolo', 'Notas']
    for i, header in enumerate(port_headers):
        header_cells3[i].text = header
        shade_cell(header_cells3[i], 'FFC7CE')
        for paragraph in header_cells3[i].paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
    
    port_data = [
        ('Portal Web', '3000', 'HTTP', 'Frontend Vue.js'),
        ('Omeka', '8081', 'HTTP', 'Repositorio'),
        ('CMS Simple', '1337', 'HTTP', 'API CMS'),
        ('Nginx', '80, 443', 'HTTP/HTTPS', 'Reverse proxy'),
        ('MySQL', '3306', 'TCP', 'Interno (no expuesto)'),
        ('Redis', '6379', 'TCP', 'Caché (opcional)')
    ]
    
    for i, (service, port, protocol, notes) in enumerate(port_data, 1):
        row_cells = table3.rows[i].cells
        row_cells[0].text = service
        row_cells[1].text = port
        row_cells[2].text = protocol
        row_cells[3].text = notes
    
    doc.add_page_break()
    
    # 7. INSTALACIÓN
    add_heading_style(doc, '7. Guía de Instalación', 1)
    
    add_heading_style(doc, '7.1 Instalación en Windows', 2)
    
    doc.add_heading('Paso 1: Clonar el repositorio', level=3)
    code_block = doc.add_paragraph('cd C:\\\\Users\\\\TuUsuario\\\\Documents', style='No Spacing')
    code_block.paragraph_format.left_indent = Inches(0.5)
    code_block.runs[0].font.name = 'Courier New'
    code_block.runs[0].font.size = Pt(10)
    
    code_block2 = doc.add_paragraph('git clone https://github.com/tu-usuario/portalweb.git', style='No Spacing')
    code_block2.paragraph_format.left_indent = Inches(0.5)
    code_block2.runs[0].font.name = 'Courier New'
    code_block2.runs[0].font.size = Pt(10)
    
    code_block3 = doc.add_paragraph('cd portalweb', style='No Spacing')
    code_block3.paragraph_format.left_indent = Inches(0.5)
    code_block3.runs[0].font.name = 'Courier New'
    code_block3.runs[0].font.size = Pt(10)
    
    doc.add_heading('Paso 2: Instalar Docker Desktop', level=3)
    doc.add_paragraph('Descargar desde: https://www.docker.com/products/docker-desktop')
    doc.add_paragraph('Ejecutar instalador')
    doc.add_paragraph('Reiniciar Windows')
    
    doc.add_heading('Paso 3: Iniciar servicios', level=3)
    commands = [
        'docker-compose build',
        'docker-compose up -d',
        'docker-compose ps'
    ]
    for cmd in commands:
        p = doc.add_paragraph(cmd, style='No Spacing')
        p.paragraph_format.left_indent = Inches(0.5)
        p.runs[0].font.name = 'Courier New'
        p.runs[0].font.size = Pt(10)
    
    doc.add_heading('Paso 4: Acceder a la aplicación', level=3)
    doc.add_paragraph('Portal: http://localhost:3000', style='List Bullet')
    doc.add_paragraph('Omeka: http://localhost:8081', style='List Bullet')
    doc.add_paragraph('CMS API: http://localhost:1337/api', style='List Bullet')
    
    add_heading_style(doc, '7.2 Instalación en macOS', 2)
    
    doc.add_paragraph('Los pasos son similares a Windows, pero usando Homebrew:')
    
    cmd_mac = doc.add_paragraph('brew install --cask docker', style='No Spacing')
    cmd_mac.paragraph_format.left_indent = Inches(0.5)
    cmd_mac.runs[0].font.name = 'Courier New'
    cmd_mac.runs[0].font.size = Pt(10)
    
    doc.add_paragraph('Luego seguir los pasos 1, 3 y 4 de Windows')
    
    add_heading_style(doc, '7.3 Instalación en Linux', 2)
    
    doc.add_paragraph('Para Ubuntu/Debian:')
    
    linux_commands = [
        'sudo apt update && sudo apt upgrade -y',
        'sudo apt install docker.io docker-compose git -y',
        'sudo usermod -aG docker $USER',
        'newgrp docker'
    ]
    
    for cmd in linux_commands:
        p = doc.add_paragraph(cmd, style='No Spacing')
        p.paragraph_format.left_indent = Inches(0.5)
        p.runs[0].font.name = 'Courier New'
        p.runs[0].font.size = Pt(9)
    
    doc.add_paragraph('Luego clonar el repositorio e iniciar los servicios con los comandos de Windows.')
    
    doc.add_page_break()
    
    # 8. CONFIGURACIÓN Y DESPLIEGUE
    add_heading_style(doc, '8. Configuración y Despliegue', 1)
    
    add_heading_style(doc, '8.1 Variables de Entorno', 2)
    
    doc.add_paragraph('Crear archivo .env.local en la raíz del proyecto:')
    
    env_content = doc.add_paragraph(
        'VITE_OMEKA_API_URL=http://localhost:8081/api\n'
        'VITE_CMS_API_URL=http://localhost:1337/api\n'
        'VITE_API_TIMEOUT=30000',
        style='No Spacing'
    )
    env_content.paragraph_format.left_indent = Inches(0.5)
    env_content.runs[0].font.name = 'Courier New'
    env_content.runs[0].font.size = Pt(10)
    
    add_heading_style(doc, '8.2 Volúmenes Docker', 2)
    
    doc.add_paragraph('Los datos persisten en volúmenes Docker:')
    
    volumes = [
        'portalweb_omeka_db: MySQL data (Omeka)',
        'portalweb_omeka_files: Archivos Omeka',
        'portalweb_cms_data: data.json (CMS)',
        'portalweb_redis_data: Redis cache'
    ]
    
    for vol in volumes:
        doc.add_paragraph(vol, style='List Bullet')
    
    add_heading_style(doc, '8.3 Comandos Útiles', 2)
    
    useful_cmds = [
        ('Ver estado de servicios', 'docker-compose ps'),
        ('Ver logs', 'docker-compose logs -f'),
        ('Reconstruir servicio', 'docker-compose build --no-cache cms'),
        ('Detener servicios', 'docker-compose down'),
        ('Limpiar volúmenes', 'docker-compose down -v'),
    ]
    
    for desc, cmd in useful_cmds:
        p = doc.add_paragraph()
        p.add_run(desc + ': ').bold = True
        p.add_run(cmd)
    
    add_heading_style(doc, '8.4 Despliegue en Producción', 2)
    
    prod_steps = [
        'Cambiar variables de entorno a dominio de producción',
        'Configurar SSL/TLS con certificados (Let\'s Encrypt)',
        'Configurar Nginx como reverse proxy',
        'Automizar backups de bases de datos',
        'Configurar monitoreo y alertas'
    ]
    
    for step in prod_steps:
        doc.add_paragraph(step, style='List Number')
    
    doc.add_page_break()
    
    # 9. API REFERENCE
    add_heading_style(doc, '9. API Reference', 1)
    
    add_heading_style(doc, '9.1 Búsqueda Global', 2)
    
    doc.add_paragraph('Endpoint: GET /api/search?query=término&source=all', style='No Spacing')
    
    search_params = doc.add_paragraph()
    search_params.add_run('Parámetros:\n').bold = True
    search_params.add_run('• query: Término a buscar\n• source: "omeka", "cms", o "all"\n• limit: Número de resultados (default: 10)')
    
    add_heading_style(doc, '9.2 Omeka API', 2)
    
    doc.add_paragraph('URL: http://localhost:8081/api/items', style='No Spacing')
    
    doc.add_paragraph('Endpoints principales:', style='List Bullet')
    doc.add_paragraph('GET /items: Listar recursos', style='List Bullet 2')
    doc.add_paragraph('GET /items/:id: Obtener recurso específico', style='List Bullet 2')
    doc.add_paragraph('GET /items?search=término: Buscar recursos', style='List Bullet 2')
    
    add_heading_style(doc, '9.3 CMS Simple API', 2)
    
    doc.add_paragraph('URL: http://localhost:1337/api', style='No Spacing')
    
    cms_endpoints = [
        ('GET /articulos', 'Listar artículos'),
        ('GET /articulos/:id', 'Obtener artículo'),
        ('POST /articulos', 'Crear artículo'),
        ('PUT /articulos/:id', 'Actualizar artículo'),
        ('DELETE /articulos/:id', 'Eliminar artículo'),
        ('GET /search?query=...', 'Buscar artículos')
    ]
    
    for endpoint, desc in cms_endpoints:
        p = doc.add_paragraph()
        p.add_run(endpoint).bold = True
        p.add_run(': ' + desc)
    
    doc.add_page_break()
    
    # 10. ESTRUCTURA DEL CÓDIGO
    add_heading_style(doc, '10. Estructura del Código', 1)
    
    add_heading_style(doc, '10.1 Árbol de Directorios', 2)
    
    tree_structure = """portalweb/
├── src/
│   ├── api/index.js
│   ├── components/
│   │   ├── Header.vue
│   │   ├── Footer.vue
│   │   └── Navigation.vue
│   ├── stores/resources.js
│   ├── views/
│   │   ├── Home.vue
│   │   ├── Search.vue
│   │   ├── Resources.vue
│   │   ├── ResourceDetail.vue
│   │   └── NotFound.vue
│   ├── App.vue
│   ├── main.js
│   └── router.js
├── cms-simple/
│   ├── server.js
│   ├── data.json
│   ├── package.json
│   ├── Dockerfile
│   └── init-data.sh
├── public/
├── docs/
├── docker-compose.yml
├── Dockerfile
├── package.json
└── README.md"""
    
    tree_para = doc.add_paragraph(tree_structure)
    tree_para.runs[0].font.name = 'Courier New'
    tree_para.runs[0].font.size = Pt(9)
    
    add_heading_style(doc, '10.2 Archivos Clave', 2)
    
    key_files = [
        ('src/api/index.js', 'Integración con APIs externas (Omeka y CMS)'),
        ('src/stores/resources.js', 'Pinia store para gestión de estado'),
        ('cms-simple/server.js', 'Servidor Express con endpoints REST'),
        ('cms-simple/data.json', 'Base de datos JSON con artículos'),
        ('docker-compose.yml', 'Orquestación de servicios Docker'),
        ('.env.local', 'Variables de entorno locales'),
    ]
    
    for filename, desc in key_files:
        p = doc.add_paragraph()
        p.add_run(filename).bold = True
        p.add_run(': ' + desc)
    
    doc.add_page_break()
    
    # 11. MANTENIMIENTO
    add_heading_style(doc, '11. Mantenimiento y Troubleshooting', 1)
    
    add_heading_style(doc, '11.1 Problemas Comunes y Soluciones', 2)
    
    problems = [
        ('Servicios no inician', 'docker-compose down -v && docker-compose build --no-cache && docker-compose up -d'),
        ('Puerto ya está en uso', 'Cambiar puerto en docker-compose.yml o liberar puerto: taskkill /PID <PID> /F'),
        ('Sin conexión a Omeka', 'docker-compose logs omeka y docker-compose restart omeka'),
        ('Acentos se ven mal en terminal', 'Esto es normal. Verificar datos en navegador.'),
    ]
    
    for problem, solution in problems:
        doc.add_heading(problem, level=3)
        p = doc.add_paragraph()
        p.add_run('Solución: ').bold = True
        p.add_run(solution)
    
    add_heading_style(doc, '11.2 Monitoreo y Logs', 2)
    
    doc.add_paragraph('Ver todos los logs:', style='List Number')
    log_cmd = doc.add_paragraph('docker-compose logs -f', style='No Spacing')
    log_cmd.paragraph_format.left_indent = Inches(0.5)
    log_cmd.runs[0].font.name = 'Courier New'
    
    doc.add_paragraph('Ver logs de servicio específico:', style='List Number')
    service_log = doc.add_paragraph('docker-compose logs -f cms', style='No Spacing')
    service_log.paragraph_format.left_indent = Inches(0.5)
    service_log.runs[0].font.name = 'Courier New'
    
    doc.add_paragraph('Buscar errores:', style='List Number')
    error_log = doc.add_paragraph('docker-compose logs | findstr ERROR', style='No Spacing')
    error_log.paragraph_format.left_indent = Inches(0.5)
    error_log.runs[0].font.name = 'Courier New'
    
    add_heading_style(doc, '11.3 Backups', 2)
    
    doc.add_paragraph('Backup de Omeka (MySQL):', style='List Number')
    backup_cmd = doc.add_paragraph('docker-compose exec omeka-db mysqldump -u root -p omeka > backup.sql', style='No Spacing')
    backup_cmd.paragraph_format.left_indent = Inches(0.5)
    backup_cmd.runs[0].font.name = 'Courier New'
    
    doc.add_paragraph('Backup de CMS (JSON):', style='List Number')
    cms_backup = doc.add_paragraph('docker cp iis-cms:/app/data.json ./data-backup.json', style='No Spacing')
    cms_backup.paragraph_format.left_indent = Inches(0.5)
    cms_backup.runs[0].font.name = 'Courier New'
    
    add_heading_style(doc, '11.4 Actualizaciones', 2)
    
    update_steps = [
        'Obtener últimos cambios: git pull origin main',
        'Reconstruir servicios: docker-compose build --no-cache',
        'Reiniciar: docker-compose down && docker-compose up -d'
    ]
    
    for step in update_steps:
        doc.add_paragraph(step, style='List Number')
    
    doc.add_page_break()
    
    # CONCLUSIÓN
    add_heading_style(doc, 'Conclusión', 1)
    
    conclusion = """El Portal Web del Instituto de Investigaciones Sociales proporciona una plataforma integrada y moderna para la gestión y consulta de recursos digitales. Con Docker y Docker Compose, el despliegue es simple y reproducible en cualquier entorno.

Las características principales incluyen:
• Búsqueda unificada en múltiples fuentes
• Interfaz moderna y responsiva
• API REST para integración
• Despliegue simple con Docker
• Mantenimiento y escalabilidad

El sistema está diseñado para ser:
• Flexible: Fácil de extender y personalizar
• Seguro: CORS habilitado, sanitización de HTML
• Performante: Caché con Redis, búsqueda optimizada
• Confiable: Volúmenes Docker para persistencia

Para soporte y actualizaciones, consulte el repositorio Git:
https://github.com/tu-usuario/portalweb"""
    
    doc.add_paragraph(conclusion)
    
    # Información final
    doc.add_page_break()
    add_heading_style(doc, 'Información del Documento', 1)
    
    final_info = [
        f'Versión: 1.0',
        f'Última actualización: Diciembre 2025',
        f'Licencia: MIT',
        f'Autor: Equipo de Desarrollo IIS',
        f'Institución: Instituto de Investigaciones Sociales, UNAM'
    ]
    
    for info in final_info:
        doc.add_paragraph(info)
    
    return doc

def main():
    """Función principal"""
    print("📄 Generando documento Word...")
    
    doc = create_documentation()
    
    output_path = r'd:\Usuarios\DEVazquezC\Documents\ICAT\IIS\portalweb\DOCUMENTACION_TECNICA.docx'
    doc.save(output_path)
    
    print(f"✅ Documento creado exitosamente: {output_path}")
    print(f"📊 Número de párrafos: {len(doc.paragraphs)}")
    print(f"📚 Número de tablas: {len(doc.tables)}")

if __name__ == '__main__':
    main()
